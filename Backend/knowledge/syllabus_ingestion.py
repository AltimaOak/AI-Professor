import os
import tempfile
from pathlib import Path

import fitz  # PyMuPDF
import docx
from PIL import Image
import pytesseract

from llama_index.core import VectorStoreIndex, StorageContext, Document
from llama_index.vector_stores.qdrant import QdrantVectorStore

from config import Config
from core.qdrant_client import get_qdrant_client

# Uncomment on Windows if tesseract is not on PATH:
# pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"


# ── Text extractors ──────────────────────────────────────────────────────────

def _extract_pdf(path: str) -> str:
    doc = fitz.open(path)
    return " ".join(page.get_text() for page in doc)

def _extract_txt(path: str) -> str:
    with open(path, "r", encoding="utf-8") as f:
        return f.read()

def _extract_docx(path: str) -> str:
    doc = docx.Document(path)
    return " ".join(para.text for para in doc.paragraphs if para.text.strip())

def _extract_image(path: str) -> str:
    image = Image.open(path)
    return pytesseract.image_to_string(image)

def _extract_text(path: str, filename: str) -> str:
    ext = Path(filename).suffix.lower()
    extractors = {
        ".pdf":  _extract_pdf,
        ".txt":  _extract_txt,
        ".docx": _extract_docx,
        ".png":  _extract_image,
        ".jpg":  _extract_image,
        ".jpeg": _extract_image,
    }
    if ext not in extractors:
        raise ValueError(f"Unsupported file type: {ext}")
    return extractors[ext](path)


# ── Main ingestion function ──────────────────────────────────────────────────

def ingest_file(file_bytes: bytes, filename: str, student_id: str) -> dict:
    """
    Parses file → chunks → embeds → stores in student's Qdrant collection.
    Each student gets their own isolated collection: ai_professor_{student_id}
    """
    suffix = Path(filename).suffix

    # Write to temp file for extraction
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        tmp.write(file_bytes)
        tmp_path = tmp.name

    try:
        raw_text = _extract_text(tmp_path, filename)
        clean_text = " ".join(raw_text.replace("\n", " ").split())

        if not clean_text.strip():
            raise ValueError("No text could be extracted from this file.")

        document = Document(
            text=clean_text,
            metadata={
                "filename": filename,
                "student_id": student_id,
            },
        )

        collection_name = f"{Config.QDRANT_COLLECTION}_{student_id}"
        client = get_qdrant_client()

        vector_store = QdrantVectorStore(
            client=client,
            collection_name=collection_name,
        )
        storage_context = StorageContext.from_defaults(vector_store=vector_store)

        # LlamaIndex handles chunking + embedding + storing automatically
        VectorStoreIndex.from_documents(
            [document],
            storage_context=storage_context,
            show_progress=False,
        )

        return {
            "status": "success",
            "filename": filename,
            "collection": collection_name,
        }

    finally:
        os.unlink(tmp_path)